"""
docx_processor.py
-----------------
DOCX reading and writing with run-level offset mapping.

Key design:
  - Every paragraph (body, table cell, header, footer) is treated as one logical text unit.
  - All runs in the paragraph are concatenated into a single string (full_text).
  - A run_map records the character-offset range of each run in full_text.
  - Presidio is run on full_text.
  - Detected PII offsets are mapped back through run_map to individual runs.
  - Redaction modifies runs at character level while preserving formatting on untouched fragments.
"""

from __future__ import annotations

import copy
import logging
from dataclasses import dataclass, field
from typing import Dict, Generator, List, Optional, Tuple

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class RunMapEntry:
    """Maps a slice of full_text back to the originating run."""
    run_index: int
    run_start: int   # inclusive offset in full_text
    run_end: int     # exclusive offset in full_text
    run_obj: Run


@dataclass
class TextUnit:
    """A single redactable content unit from a DOCX document."""
    unit_id: str
    full_text: str
    paragraph: Paragraph
    run_map: List[RunMapEntry] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------

def _build_text_unit(paragraph: Paragraph, unit_id: str) -> TextUnit:
    """Concatenate all runs of a paragraph into full_text and build run_map."""
    parts: List[str] = []
    run_map: List[RunMapEntry] = []
    offset = 0
    for idx, run in enumerate(paragraph.runs):
        text = run.text or ""
        run_map.append(RunMapEntry(
            run_index=idx,
            run_start=offset,
            run_end=offset + len(text),
            run_obj=run,
        ))
        parts.append(text)
        offset += len(text)
    return TextUnit(
        unit_id=unit_id,
        full_text="".join(parts),
        paragraph=paragraph,
        run_map=run_map,
    )


def _iter_paragraphs_in_table(table, table_prefix: str) -> Generator[TextUnit, None, None]:
    """Recursively yield TextUnits for all cells (including nested tables)."""
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            for p_idx, para in enumerate(cell.paragraphs):
                unit_id = f"{table_prefix}-r{r_idx}-c{c_idx}-para-{p_idx}"
                tu = _build_text_unit(para, unit_id)
                if tu.full_text.strip():
                    yield tu
            # Recurse into nested tables
            for nested_idx, nested_table in enumerate(cell.tables):
                nested_prefix = f"{table_prefix}-nested-{nested_idx}"
                yield from _iter_paragraphs_in_table(nested_table, nested_prefix)


def extract_text_units(doc: Document) -> Generator[TextUnit, None, None]:
    """
    Yield every redactable TextUnit from the document:
      - Body paragraphs
      - Tables (all rows/cells, nested)
      - Per-section headers
      - Per-section footers
    """
    # --- Body paragraphs ---
    for p_idx, para in enumerate(doc.paragraphs):
        unit_id = f"body-para-{p_idx}"
        tu = _build_text_unit(para, unit_id)
        if tu.full_text.strip():
            yield tu

    # --- Tables ---
    for t_idx, table in enumerate(doc.tables):
        table_prefix = f"table-{t_idx}"
        yield from _iter_paragraphs_in_table(table, table_prefix)

    # --- Per-section headers and footers ---
    for s_idx, section in enumerate(doc.sections):
        # Header
        header = section.header
        if header is not None:
            for p_idx, para in enumerate(header.paragraphs):
                unit_id = f"section-{s_idx}-header-para-{p_idx}"
                tu = _build_text_unit(para, unit_id)
                if tu.full_text.strip():
                    yield tu
            # Tables inside header
            for t_idx_h, tbl in enumerate(header.tables):
                tbl_prefix = f"section-{s_idx}-header-table-{t_idx_h}"
                yield from _iter_paragraphs_in_table(tbl, tbl_prefix)

        # Footer
        footer = section.footer
        if footer is not None:
            for p_idx, para in enumerate(footer.paragraphs):
                unit_id = f"section-{s_idx}-footer-para-{p_idx}"
                tu = _build_text_unit(para, unit_id)
                if tu.full_text.strip():
                    yield tu
            # Tables inside footer
            for t_idx_f, tbl in enumerate(footer.tables):
                tbl_prefix = f"section-{s_idx}-footer-table-{t_idx_f}"
                yield from _iter_paragraphs_in_table(tbl, tbl_prefix)


# ---------------------------------------------------------------------------
# Redaction write-back
# ---------------------------------------------------------------------------

def apply_redactions(unit: TextUnit, replacements: List[Tuple[int, int, str]]) -> None:
    """
    Apply (start, end, replacement_text) tuples to a TextUnit in-place.
    Operates at the run level to preserve formatting.

    replacements must be sorted in DESCENDING order by start offset so that
    earlier replacements don't shift offsets of later ones.
    """
    if not replacements:
        return

    # Build the full text with replacements applied to locate run boundaries
    # We use a "character annotation" approach:
    # For each character position in full_text, record its run_map entry.
    full = unit.full_text
    char_to_run: List[int] = []
    for entry in unit.run_map:
        char_to_run.extend([entry.run_index] * (entry.run_end - entry.run_start))

    # Process replacements from right to left (already sorted desc by caller)
    for (pii_start, pii_end, replacement) in replacements:
        if pii_start >= pii_end or pii_start >= len(full):
            continue
        pii_end = min(pii_end, len(full))

        # Identify which runs are touched
        touched_run_indices = set(char_to_run[pii_start:pii_end])
        touched_entries = [e for e in unit.run_map if e.run_index in touched_run_indices]
        if not touched_entries:
            continue

        # Sort touched entries by run_index
        touched_entries.sort(key=lambda e: e.run_index)
        first_entry = touched_entries[0]
        last_entry = touched_entries[-1]

        # Compute how much of the first run to keep before the PII span
        prefix_in_first = full[first_entry.run_start:pii_start]
        # Compute how much of the last run to keep after the PII span
        suffix_in_last = full[pii_end:last_entry.run_end]

        # Modify first run: prefix + replacement
        first_run = first_entry.run_obj
        first_run.text = prefix_in_first + replacement

        # Clear intermediate runs (between first and last)
        for entry in touched_entries[1:]:
            entry.run_obj.text = ""

        # Add suffix to last run (if different from first)
        if last_entry.run_index != first_entry.run_index:
            last_entry.run_obj.text = suffix_in_last

        # Update the in-memory full_text and char_to_run to keep subsequent
        # replacements consistent (right-to-left order makes this safe)
        new_full = full[:pii_start] + replacement + full[pii_end:]
        # Rebuild char_to_run for remaining (earlier) replacements
        delta = len(replacement) - (pii_end - pii_start)
        # We only need positions 0..pii_start, which are unchanged
        full = new_full
        # Truncate char_to_run to the positions before pii_start (unchanged)
        char_to_run = char_to_run[:pii_start] + [first_entry.run_index] * len(replacement) + [
            char_to_run[i] for i in range(pii_end, len(char_to_run))
        ]
        # Update run_map entries for affected and subsequent runs
        for entry in unit.run_map:
            if entry.run_index < first_entry.run_index:
                pass  # unchanged
            elif entry.run_index == first_entry.run_index:
                entry.run_end = entry.run_start + len(first_run.text)
            elif entry.run_index in {e.run_index for e in touched_entries[1:]}:
                entry.run_start = first_entry.run_end
                entry.run_end = first_entry.run_end
            else:
                entry.run_start += delta
                entry.run_end += delta
        unit.full_text = full


def load_document(input_path: str) -> Document:
    """Load a DOCX and return the Document object."""
    try:
        doc = Document(input_path)
        logger.info("Loaded document: %s", input_path)
        return doc
    except Exception as exc:
        raise RuntimeError(f"Failed to load DOCX '{input_path}': {exc}") from exc


def save_document(doc: Document, output_path: str) -> None:
    """Save the document to output_path (never overwrites input)."""
    try:
        doc.save(output_path)
        logger.info("Saved redacted document: %s", output_path)
    except Exception as exc:
        raise RuntimeError(f"Failed to save DOCX '{output_path}': {exc}") from exc
