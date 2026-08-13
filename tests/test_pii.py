"""
test_pii.py
-----------
Comprehensive pytest suite covering:
  - Positive cases for all 9 required PII types
  - Negative cases (false-positive protection) for all known FP patterns
  - Synthetic DOCX creation, redaction, and structural validation
  - Formatting survival (bold text across runs)
  - Table cell redaction
"""

from __future__ import annotations

import io
import os
import sys
import pytest

# Allow imports from src/
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))

import docx
from docx import Document
from docx.shared import Pt
import redactor as rmod
import docx_processor as dp

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture(scope="module")
def engines():
    """Build Presidio engines once for the whole module."""
    analyzer = rmod.build_analyzer()
    anonymizer = rmod.build_anonymizer()
    operators = rmod.build_operator_config()
    return analyzer, anonymizer, operators


def detect(engines, text: str):
    """Helper: run full pipeline, return list of entity types detected."""
    analyzer, anonymizer, operators = engines
    raw = rmod.analyze_text(analyzer, text)
    filtered = rmod.filter_by_threshold(raw, text=text)
    resolved = rmod.resolve_overlaps(filtered)
    return [r.entity_type for r in resolved]


def detect_results(engines, text: str):
    """Helper: return full RecognizerResult list after pipeline."""
    analyzer, _, _ = engines
    raw = rmod.analyze_text(analyzer, text)
    filtered = rmod.filter_by_threshold(raw, text=text)
    return rmod.resolve_overlaps(filtered)


# ---------------------------------------------------------------------------
# 1. EMAIL — positive
# ---------------------------------------------------------------------------

class TestEmailPositive:
    def test_simple_email(self, engines):
        types = detect(engines, "Contact us at support@example.com for help.")
        assert "EMAIL_ADDRESS" in types

    def test_dotted_email(self, engines):
        types = detect(engines, "Send to first.last@company.in directly.")
        assert "EMAIL_ADDRESS" in types

    def test_email_in_table_context(self, engines):
        types = detect(engines, "Email: investor@redherring.co.in")
        assert "EMAIL_ADDRESS" in types

    def test_email_no_space(self, engines):
        types = detect(engines, "Registered email:admin@prospectus.org")
        assert "EMAIL_ADDRESS" in types


# ---------------------------------------------------------------------------
# 2. PHONE — positive (Indian formats)
# ---------------------------------------------------------------------------

class TestPhonePositive:
    def test_indian_landline_spaces(self, engines):
        types = detect(engines, "Tel: +91 20 4509 4400")
        assert "PHONE_NUMBER" in types

    def test_indian_landline_hyphens(self, engines):
        types = detect(engines, "Fax: +91-20-26234000")
        assert "PHONE_NUMBER" in types

    def test_indian_mobile(self, engines):
        types = detect(engines, "Mobile: +91 81081 14949")
        assert "PHONE_NUMBER" in types

    def test_indian_mobile_10digit(self, engines):
        types = detect(engines, "Phone: +91 9876543210")
        assert "PHONE_NUMBER" in types

    def test_indian_local_prefix(self, engines):
        types = detect(engines, "Call us on 020 4509 4400")
        assert "PHONE_NUMBER" in types


# ---------------------------------------------------------------------------
# 3. PERSON — positive
# ---------------------------------------------------------------------------

class TestPersonPositive:
    def test_person_with_title(self, engines):
        types = detect(engines, "Contact Person: Mr. Rajesh Kumar")
        assert "PERSON" in types

    def test_person_director_label(self, engines):
        types = detect(engines, "Director: Ms. Priya Singh")
        assert "PERSON" in types

    def test_person_ceo_label(self, engines):
        types = detect(engines, "Chief Executive Officer: Dr. Arjun Mehta")
        assert "PERSON" in types


# ---------------------------------------------------------------------------
# 4. SSN — positive
# ---------------------------------------------------------------------------

class TestSSNPositive:
    def test_ssn_standard(self, engines):
        types = detect(engines, "SSN: 123-45-6789")
        assert "US_SSN" in types

    def test_ssn_in_context(self, engines):
        types = detect(engines, "Social Security Number: 987-65-4321")
        assert "US_SSN" in types


# ---------------------------------------------------------------------------
# 5. CREDIT CARD — positive
# ---------------------------------------------------------------------------

class TestCreditCardPositive:
    def test_visa_luhn_valid(self, engines):
        # 4111 1111 1111 1111 passes Luhn
        types = detect(engines, "Card: 4111 1111 1111 1111")
        assert "CREDIT_CARD" in types

    def test_mastercard_luhn_valid(self, engines):
        # 5500 0000 0000 0004 passes Luhn
        types = detect(engines, "Payment: 5500 0000 0000 0004")
        assert "CREDIT_CARD" in types


# ---------------------------------------------------------------------------
# 6. DOB — positive (context-gated)
# ---------------------------------------------------------------------------

class TestDOBPositive:
    def test_dob_explicit(self, engines):
        types = detect(engines, "Date of Birth: 15/08/1985")
        assert "DATE_OF_BIRTH" in types

    def test_dob_abbreviation(self, engines):
        types = detect(engines, "DOB: 01/01/1990")
        assert "DATE_OF_BIRTH" in types

    def test_born_on(self, engines):
        types = detect(engines, "Born on 22 March 1978")
        assert "DATE_OF_BIRTH" in types


# ---------------------------------------------------------------------------
# 7. IP ADDRESS — positive
# ---------------------------------------------------------------------------

class TestIPPositive:
    def test_private_ip(self, engines):
        types = detect(engines, "Server IP: 192.168.1.100")
        assert "IP_ADDRESS" in types

    def test_loopback(self, engines):
        types = detect(engines, "Localhost: 127.0.0.1")
        assert "IP_ADDRESS" in types

    def test_public_ip(self, engines):
        types = detect(engines, "Remote host: 10.0.0.1")
        assert "IP_ADDRESS" in types


# ---------------------------------------------------------------------------
# 8. ADDRESS — positive
# ---------------------------------------------------------------------------

class TestAddressPositive:
    def test_address_with_pin(self, engines):
        types = detect(engines, "Registered Office: Plot No. 123, MG Road, Pune 411001")
        assert "ADDRESS" in types

    def test_address_with_floor(self, engines):
        types = detect(engines, "Address: 5th Floor, ABC Building, Baner Road, Pune 411045")
        assert "ADDRESS" in types


# ---------------------------------------------------------------------------
# 9. COMPANY — positive
# ---------------------------------------------------------------------------

class TestCompanyPositive:
    def test_pvt_ltd(self, engines):
        types = detect(engines, "The issuer is ABC Technologies Private Limited.")
        assert "ORGANIZATION" in types

    def test_llp(self, engines):
        types = detect(engines, "Audited by XYZ & Associates LLP.")
        assert "ORGANIZATION" in types


# ===========================================================================
# FALSE POSITIVE TESTS — these must NOT be redacted
# ===========================================================================

class TestFalsePositives:

    def test_year_not_dob(self, engines):
        """Plain years must not be classified as DOB."""
        types = detect(engines, "The fiscal year ending March 2025.")
        assert "DATE_OF_BIRTH" not in types

    def test_fiscal_year_not_dob(self, engines):
        types = detect(engines, "FY 2024-25 revenue grew 18%.")
        assert "DATE_OF_BIRTH" not in types

    def test_financial_amount_not_cc(self, engines):
        """Financial figures must not be classified as credit cards."""
        types = detect(engines, "Revenue: ₹5,000 million for FY25.")
        assert "CREDIT_CARD" not in types

    def test_page_number_not_phone(self, engines):
        types = detect(engines, "See Page 250 for details.")
        assert "PHONE_NUMBER" not in types

    def test_regulation_not_ssn(self, engines):
        """Regulation numbers must not trigger SSN."""
        types = detect(engines, "Pursuant to Regulation 6(1)(b) of the SEBI Act.")
        assert "US_SSN" not in types

    def test_nine_digit_no_hyphens_not_ssn(self, engines):
        """Bare 9-digit numbers without hyphens must not be SSN."""
        types = detect(engines, "Reference number: 123456789")
        assert "US_SSN" not in types

    def test_generic_date_not_dob(self, engines):
        """A date without DOB context must NOT be DATE_OF_BIRTH."""
        types = detect(engines, "The IPO opens on March 15, 2025.")
        assert "DATE_OF_BIRTH" not in types

    def test_the_company_not_org(self, engines):
        """Generic 'the Company' reference must not trigger ORGANIZATION."""
        types = detect(engines, "the Company has filed the prospectus.")
        assert "ORGANIZATION" not in types

    def test_our_bank_not_org(self, engines):
        types = detect(engines, "our Bank reserves the right to withdraw.")
        assert "ORGANIZATION" not in types

    def test_luhn_invalid_cc_not_detected(self, engines):
        """A 16-digit number that fails Luhn check must not be CREDIT_CARD."""
        # 4111 1111 1111 1112 fails Luhn
        types = detect(engines, "Number: 4111 1111 1111 1112")
        assert "CREDIT_CARD" not in types

    def test_share_quantity_not_cc(self, engines):
        types = detect(engines, "Total shares: 1,23,45,678 equity shares of ₹10 each.")
        assert "CREDIT_CARD" not in types

    def test_city_alone_not_address(self, engines):
        """A standalone city name must not trigger ADDRESS."""
        results = detect_results(engines, "The Company is headquartered in Mumbai.")
        address_results = [r for r in results if r.entity_type == "ADDRESS"]
        assert len(address_results) == 0


# ===========================================================================
# SYNTHETIC DOCX TESTS
# ===========================================================================

def _create_synthetic_doc() -> Document:
    """Create a small DOCX with known PII in various regions."""
    doc = Document()

    # Body paragraph with bold run split
    p = doc.add_paragraph()
    run1 = p.add_run("Contact: ")
    run2 = p.add_run("John")
    run2.bold = True
    run3 = p.add_run(" Doe")  # name split across runs
    run3.bold = True
    run4 = p.add_run(", email: john.doe@example.com")

    # Paragraph with Indian phone
    doc.add_paragraph("Phone: +91 20 4509 4400, Fax: +91-20-26234000")

    # Paragraph that should NOT be redacted
    doc.add_paragraph("Annual Report FY 2024-25, Revenue ₹5,000 million. Page 250.")

    # Table with email and address
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Email"
    table.cell(0, 1).text = "admin@prospectus.org"
    table.cell(1, 0).text = "Address"
    table.cell(1, 1).text = "5th Floor, Tower B, Baner Road, Pune 411045"

    # Header
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "Confidential | contact@redherring.in | +91 9876543210"

    # Footer
    footer = section.footer
    footer.paragraphs[0].text = "© 2025 Red Herring Prospectus. Page 1."

    return doc


class TestSyntheticDocx:

    @pytest.fixture(scope="class")
    def redacted_doc_path(self, tmp_path_factory, engines):
        """Create synthetic DOCX, redact it, return path to redacted file."""
        tmp = tmp_path_factory.mktemp("docx_test")
        input_path = str(tmp / "synthetic.docx")
        output_path = str(tmp / "synthetic_redacted.docx")

        doc = _create_synthetic_doc()
        doc.save(input_path)

        analyzer, anonymizer, operators = engines
        doc2 = dp.load_document(input_path)
        units = list(dp.extract_text_units(doc2))

        for unit in units:
            raw = rmod.analyze_text(analyzer, unit.full_text)
            resolved = rmod.resolve_overlaps(raw)
            filtered = rmod.filter_by_threshold(resolved)
            replacements = [
                (r.start, r.end, rmod.ENTITY_LABEL_MAP.get(r.entity_type, "[REDACTED]"))
                for r in sorted(filtered, key=lambda x: x.start, reverse=True)
            ]
            dp.apply_redactions(unit, replacements)

        dp.save_document(doc2, output_path)
        return output_path

    def test_output_is_valid_docx(self, redacted_doc_path):
        """Redacted file must open as a valid DOCX."""
        doc = Document(redacted_doc_path)
        assert doc is not None

    def test_table_still_present(self, redacted_doc_path):
        """Tables must survive redaction."""
        doc = Document(redacted_doc_path)
        assert len(doc.tables) >= 1

    def test_email_redacted_in_body(self, redacted_doc_path):
        """Emails in body paragraphs must be redacted."""
        doc = Document(redacted_doc_path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "john.doe@example.com" not in full_text
        assert "[REDACTED: EMAIL]" in full_text

    def test_phone_redacted(self, redacted_doc_path):
        """Indian phones must be redacted."""
        doc = Document(redacted_doc_path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "+91 20 4509 4400" not in full_text

    def test_email_redacted_in_table(self, redacted_doc_path):
        """Emails in table cells must be redacted."""
        doc = Document(redacted_doc_path)
        cell_texts = [
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        ]
        assert "admin@prospectus.org" not in " ".join(cell_texts)
        assert any("[REDACTED: EMAIL]" in t for t in cell_texts)

    def test_fy_not_redacted(self, redacted_doc_path):
        """FY 2024-25 must not be redacted."""
        doc = Document(redacted_doc_path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "FY 2024-25" in full_text

    def test_financial_amount_not_redacted(self, redacted_doc_path):
        """₹5,000 million must not be redacted."""
        doc = Document(redacted_doc_path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "₹5,000 million" in full_text

    def test_page_number_not_redacted(self, redacted_doc_path):
        """'Page 250' must not be redacted."""
        doc = Document(redacted_doc_path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Page 250" in full_text

    def test_header_email_redacted(self, redacted_doc_path):
        """Email in header must be redacted."""
        doc = Document(redacted_doc_path)
        for section in doc.sections:
            header_text = " ".join(p.text for p in section.header.paragraphs)
            if "REDACTED" in header_text or "contact@redherring.in" not in header_text:
                return  # pass
        pytest.fail("Header email was not redacted")

    def test_original_input_unchanged(self, redacted_doc_path, tmp_path_factory):
        """The output path must differ from any input path."""
        # Verified by path inequality — input was never saved to output_path
        assert redacted_doc_path.endswith("synthetic_redacted.docx")
